#!/usr/bin/env python3
"""
Document Processing Pipeline

Process different file formats in the current directory,
convert them to Markdown, and combine them into a single output file.
"""

import argparse
import datetime
import glob
import logging
import os
import sys
import time
import urllib.parse
from typing import Dict, List, Optional, Tuple, Any

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    datefmt="%H:%M:%S"
)
logger = logging.getLogger(__name__)

# Try importing optional dependencies
DEPENDENCIES = {
    "pdf": False,
    "docx": False,
    "html": False,
    "csv": False,
    "openai": False,
    "requests": False
}

try:
    import pdfplumber
    DEPENDENCIES["pdf"] = True
except ImportError:
    logger.warning("pdfplumber not installed. PDF processing will be limited.")

try:
    from docx import Document
    DEPENDENCIES["docx"] = True
except ImportError:
    logger.warning("python-docx not installed. DOCX processing will be limited.")

try:
    from markdownify import markdownify as md
    from bs4 import BeautifulSoup
    DEPENDENCIES["html"] = True
except ImportError:
    logger.warning("markdownify/bs4 not installed. HTML processing will be limited.")

try:
    import pandas as pd
    DEPENDENCIES["csv"] = True
except ImportError:
    logger.warning("pandas not installed. CSV processing will be limited.")

try:
    import openai
    DEPENDENCIES["openai"] = True
except ImportError:
    logger.warning("openai not installed. Summary generation will be disabled.")

try:
    import requests
    DEPENDENCIES["requests"] = True
except ImportError:
    logger.warning("requests not installed. Web fetching will be disabled.")


def parse_arguments() -> argparse.Namespace:
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Process documents and convert them to a combined Markdown file."
    )

    # Add optional positional argument for input file
    parser.add_argument(
        "input_path",
        nargs="?",
        help="Path to a file or directory to process"
    )

    # Create a mutually exclusive group for input sources (only used if no positional argument)
    input_group = parser.add_mutually_exclusive_group(required=False)
    input_group.add_argument(
        "--file",
        help="Single file to process"
    )
    input_group.add_argument(
        "--directory",
        help="Directory containing files to process"
    )
    input_group.add_argument(
        "--url",
        help="URL to download and convert to Markdown"
    )

    parser.add_argument(
        "--metadata",
        nargs="+",
        help="Metadata as key=value pairs (e.g., --metadata author='Jane Doe' date='2023-01-01')"
    )
    parser.add_argument(
        "--tags",
        nargs="+",
        help="Tags to include in the document (e.g., --tags research documentation)"
    )
    parser.add_argument(
        "--summary",
        action="store_true",
        help="Generate summaries using OpenAI"
    )
    parser.add_argument(
        "--model",
        default="gpt-3.5-turbo",
        help="OpenAI model to use for summarization (default: gpt-3.5-turbo)"
    )
    parser.add_argument(
        "--length",
        type=int,
        default=150,
        help="Maximum length of summaries in tokens (default: 150)"
    )
    parser.add_argument(
        "--output",
        help="Output filename (defaults to input filename with .md extension for files, or combined_output.md for directories)"
    )

    return parser.parse_args()


def find_files(input_dir: str = ".") -> Dict[str, List[str]]:
    """
    Find files in the specified directory by extension.
    
    Args:
        input_dir: Directory to search for files
        
    Returns:
        Dict mapping file extensions to lists of matching filenames
    """
    # Convert relative to absolute path if needed
    input_dir = os.path.abspath(input_dir)
    
    if not os.path.isdir(input_dir):
        logger.error(f"Input directory does not exist: {input_dir}")
        return {ext: [] for ext in ["pdf", "docx", "txt", "html", "csv", "md"]}
    
    logger.info(f"Searching for files in: {input_dir}")
    
    # Build search patterns with directory path
    extensions = {
        "pdf": glob.glob(os.path.join(input_dir, "*.pdf")),
        "docx": glob.glob(os.path.join(input_dir, "*.docx")),
        "txt": glob.glob(os.path.join(input_dir, "*.txt")),
        "html": glob.glob(os.path.join(input_dir, "*.html")) + glob.glob(os.path.join(input_dir, "*.htm")),
        "csv": glob.glob(os.path.join(input_dir, "*.csv")),
        "md": glob.glob(os.path.join(input_dir, "*.md"))
    }
    
    # Log what we found
    total_files = sum(len(files) for files in extensions.values())
    logger.info(f"Found {total_files} files to process")
    for ext, files in extensions.items():
        if files:
            logger.info(f"  {len(files)} {ext.upper()} files")
    
    return extensions


def convert_pdf(filepath: str) -> str:
    """
    Convert PDF file to Markdown.
    
    Args:
        filepath: Path to the PDF file
        
    Returns:
        Markdown string representation
    """
    if not DEPENDENCIES["pdf"]:
        return f"*Error: Could not process PDF file. Please install pdfplumber.*\n\n"
    
    try:
        markdown_content = []
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    markdown_content.append(f"### Page {i+1}\n\n{text}\n\n")
                else:
                    markdown_content.append(f"### Page {i+1}\n\n*No extractable text on this page*\n\n")
        
        return "\n".join(markdown_content)
    except Exception as e:
        logger.error(f"Error processing PDF file {filepath}: {str(e)}")
        return f"*Error: Failed to process PDF file: {str(e)}*\n\n"


def convert_docx(filepath: str) -> str:
    """
    Convert DOCX file to Markdown.
    
    Args:
        filepath: Path to the DOCX file
        
    Returns:
        Markdown string representation
    """
    if not DEPENDENCIES["docx"]:
        return f"*Error: Could not process DOCX file. Please install python-docx.*\n\n"
    
    try:
        doc = Document(filepath)
        content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1])
                    content.append(f"{'#' * level} {para.text}\n\n")
                else:
                    text = para.text
                    # Basic formatting
                    if para.style.name == 'List Bullet':
                        text = f"* {text}"
                    elif para.style.name == 'List Number':
                        text = f"1. {text}"
                    content.append(f"{text}\n\n")
        
        # Process tables
        for table in doc.tables:
            table_rows = []
            # Header row
            header = " | ".join(cell.text for cell in table.rows[0].cells)
            separator = "|" + "---|" * len(table.rows[0].cells)
            table_rows.append(f"| {header} |")
            table_rows.append(separator)
            
            # Data rows
            for row in table.rows[1:]:
                row_data = " | ".join(cell.text for cell in row.cells)
                table_rows.append(f"| {row_data} |")
            
            content.append("\n".join(table_rows) + "\n\n")
        
        return "\n".join(content)
    except Exception as e:
        logger.error(f"Error processing DOCX file {filepath}: {str(e)}")
        return f"*Error: Failed to process DOCX file: {str(e)}*\n\n"


def convert_txt(filepath: str) -> str:
    """
    Convert TXT file to Markdown (minimal conversion needed).
    
    Args:
        filepath: Path to the TXT file
        
    Returns:
        Markdown string representation
    """
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        return content + "\n\n"
    except Exception as e:
        logger.error(f"Error processing TXT file {filepath}: {str(e)}")
        return f"*Error: Failed to process TXT file: {str(e)}*\n\n"


def convert_html(filepath: str) -> str:
    """
    Convert HTML file to Markdown.
    
    Args:
        filepath: Path to the HTML file
        
    Returns:
        Markdown string representation
    """
    if not DEPENDENCIES["html"]:
        return f"*Error: Could not process HTML file. Please install markdownify and beautifulsoup4.*\n\n"
    
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Convert to markdown
        markdown_content = md(str(soup))
        return markdown_content + "\n\n"
    except Exception as e:
        logger.error(f"Error processing HTML file {filepath}: {str(e)}")
        return f"*Error: Failed to process HTML file: {str(e)}*\n\n"


def convert_csv(filepath: str) -> str:
    """
    Convert CSV file to Markdown table.
    
    Args:
        filepath: Path to the CSV file
        
    Returns:
        Markdown string representation
    """
    if not DEPENDENCIES["csv"]:
        return f"*Error: Could not process CSV file. Please install pandas.*\n\n"
    
    try:
        df = pd.read_csv(filepath)
        return df.to_markdown(index=False) + "\n\n"
    except Exception as e:
        logger.error(f"Error processing CSV file {filepath}: {str(e)}")
        return f"*Error: Failed to process CSV file: {str(e)}*\n\n"


def convert_md(filepath: str) -> str:
    """
    Process Markdown file (already in target format).
    
    Args:
        filepath: Path to the Markdown file
        
    Returns:
        Markdown string representation
    """
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        return content + "\n\n"
    except Exception as e:
        logger.error(f"Error processing Markdown file {filepath}: {str(e)}")
        return f"*Error: Failed to process Markdown file: {str(e)}*\n\n"


def fetch_url(url: str) -> Tuple[str, str]:
    """
    Fetch content from a URL and convert it to markdown.
    
    Args:
        url: URL to fetch
        
    Returns:
        Tuple of (title, markdown_content)
    """
    if not DEPENDENCIES["requests"] or not DEPENDENCIES["html"]:
        error_msg = "*Error: Cannot fetch URL. Please install requests and beautifulsoup4/markdownify.*"
        return "URL Fetch Failed", error_msg
    
    try:
        # Add scheme if not present
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        # Make the request
        logger.info(f"Fetching URL: {url}")
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Parse the HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Extract title
        title = soup.title.string if soup.title else urllib.parse.urlparse(url).netloc
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Convert to markdown
        markdown_content = md(str(soup))
        
        # Add URL reference
        markdown_content = f"*Source: [{url}]({url})*\n\n{markdown_content}"
        
        return title, markdown_content
    
    except Exception as e:
        logger.error(f"Error fetching URL {url}: {str(e)}")
        return "URL Fetch Failed", f"*Error: Failed to fetch URL: {str(e)}*\n\n"


def generate_summary(text: str, model: str, max_length: int) -> str:
    """
    Generate a summary of the text using OpenAI API.
    
    Args:
        text: Text to summarize
        model: OpenAI model to use
        max_length: Maximum length of summary in tokens
        
    Returns:
        Generated summary
    """
    if not DEPENDENCIES["openai"]:
        return "*Summary generation skipped: OpenAI package not installed*\n\n"
    
    # Truncate text if too long (OpenAI has token limits)
    max_chars = 6000  # Approximate token to character ratio
    if len(text) > max_chars:
        text = text[:max_chars] + "..."
    
    try:
        response = openai.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": f"Summarize the following text in {max_length} tokens or less:"},
                {"role": "user", "content": text}
            ],
            max_tokens=max_length
        )
        
        summary = response.choices[0].message.content.strip()
        return f"**Summary:** {summary}\n\n"
    except Exception as e:
        logger.error(f"Error generating summary: {str(e)}")
        return f"*Summary generation failed: {str(e)}*\n\n"


def process_file(filepath: str, file_ext: str, enable_summary: bool, model: str, max_length: int) -> Tuple[str, str]:
    """
    Process a single file based on its extension.
    
    Args:
        filepath: Path to the file
        file_ext: File extension (without the dot)
        enable_summary: Whether to generate a summary
        model: OpenAI model to use for summarization
        max_length: Maximum summary length
        
    Returns:
        Tuple of (content, summary)
    """
    logger.info(f"Processing: {filepath}")
    
    # Convert based on file type
    if file_ext == "pdf":
        content = convert_pdf(filepath)
    elif file_ext == "docx":
        content = convert_docx(filepath)
    elif file_ext == "txt":
        content = convert_txt(filepath)
    elif file_ext == "html" or file_ext == "htm":
        content = convert_html(filepath)
    elif file_ext == "csv":
        content = convert_csv(filepath)
    elif file_ext == "md":
        content = convert_md(filepath)
    else:
        content = f"*Unsupported file type: {file_ext}*\n\n"
    
    # Generate summary if requested
    summary = ""
    if enable_summary and content and not content.startswith("*Error:"):
        logger.info(f"Generating summary for: {filepath}")
        # Add rate limiting to avoid API throttling
        time.sleep(1)
        summary = generate_summary(content, model, max_length)
    
    return content, summary


def parse_metadata(metadata_args: Optional[List[str]]) -> Dict[str, str]:
    """
    Parse metadata arguments from command line.
    
    Args:
        metadata_args: List of key=value strings
        
    Returns:
        Dictionary of metadata
    """
    metadata = {}
    if metadata_args:
        for item in metadata_args:
            if "=" in item:
                key, value = item.split("=", 1)
                metadata[key.strip()] = value.strip()
            else:
                logger.warning(f"Ignoring invalid metadata format: {item}")
    
    # Add automatic metadata
    metadata["processing_date"] = datetime.datetime.now().strftime("%Y-%m-%d")
    metadata["processor"] = "Document Processing Pipeline"
    
    return metadata


def generate_yaml_frontmatter(metadata: Dict[str, str], tags: Optional[List[str]]) -> str:
    """
    Generate YAML frontmatter from metadata and tags.
    
    Args:
        metadata: Dictionary of metadata
        tags: List of tags
        
    Returns:
        YAML frontmatter string
    """
    frontmatter = ["---"]
    
    # Add metadata
    for key, value in metadata.items():
        frontmatter.append(f"{key}: {value}")
    
    # Add tags if provided
    if tags and len(tags) > 0:
        frontmatter.append("tags:")
        for tag in tags:
            frontmatter.append(f"  - {tag}")
    
    frontmatter.append("---")
    return "\n".join(frontmatter) + "\n\n"


def process_url_to_markdown(url: str, output_file: str, metadata: Dict[str, str], 
                           tags: Optional[List[str]], enable_summary: bool, 
                           model: str, max_length: int) -> int:
    """
    Process a URL directly to a markdown file.
    
    Args:
        url: URL to process
        output_file: Output file path
        metadata: Dictionary of metadata
        tags: List of tags
        enable_summary: Whether to generate a summary
        model: OpenAI model to use
        max_length: Maximum summary length
        
    Returns:
        Exit code (0 for success, 1 for failure)
    """
    # Fetch the URL
    title, content = fetch_url(url)
    
    if content.startswith("*Error:"):
        logger.error(f"Failed to process URL: {url}")
        return 1
    
    # Update metadata with title and URL
    metadata["title"] = title
    metadata["source_url"] = url
    
    # Start generating output
    output_parts = []
    
    # Add frontmatter
    output_parts.append(generate_yaml_frontmatter(metadata, tags))
    
    # Add title
    output_parts.append(f"# {title}\n\n")
    
    # Generate summary if requested
    if enable_summary and DEPENDENCIES["openai"]:
        logger.info(f"Generating summary for URL: {url}")
        summary = generate_summary(content, model, max_length)
        output_parts.append(summary)
    
    # Add content
    output_parts.append(content)
    
    # Write to output file
    with open(output_file, "w", encoding="utf-8") as f:
        f.write("\n".join(output_parts))
    
    logger.info(f"URL content written to {output_file}")
    return 0


def main() -> int:
    """Main function."""
    args = parse_arguments()

    # Set OpenAI API key if summary is enabled
    if args.summary:
        if DEPENDENCIES["openai"]:
            openai_key = os.environ.get("OPENAI_API_KEY")
            if not openai_key:
                logger.error("OPENAI_API_KEY environment variable not set. Summaries will be disabled.")
                args.summary = False
            else:
                openai.api_key = openai_key
                logger.info(f"Using OpenAI model: {args.model}")
        else:
            args.summary = False

    # Parse metadata
    metadata = parse_metadata(args.metadata)

    # Process positional argument if provided
    if args.input_path:
        # Determine if it's a file or directory
        if os.path.isfile(args.input_path):
            # It's a file, process it as such
            file_path = args.input_path

            # Determine output filename
            output_file = args.output
            if not output_file:
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                output_file = f"{base_name}.md"
                logger.info(f"Output filename not specified, using: {output_file}")

            file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')

            # Check if file extension is supported
            if file_ext not in ["pdf", "docx", "txt", "html", "htm", "csv", "md"]:
                logger.error(f"Unsupported file type: {file_ext}")
                return 1

            # Process the single file
            logger.info(f"Processing single file: {file_path}")

            # Start generating output
            output_parts = []

            # Add frontmatter
            output_parts.append(generate_yaml_frontmatter(metadata, args.tags))

            # Add title
            output_parts.append(f"# {os.path.basename(file_path)}\n\n")

            # Process file
            content, summary = process_file(
                file_path, file_ext, args.summary, args.model, args.length
            )

            # Add summary if generated
            if summary:
                output_parts.append(summary)

            # Add content
            output_parts.append(content)

            # Write to output file
            with open(output_file, "w", encoding="utf-8") as f:
                f.write("\n".join(output_parts))

            logger.info(f"Output written to {output_file}")
            return 0

        elif os.path.isdir(args.input_path):
            # It's a directory
            directory_path = args.input_path

            # Determine output filename
            output_file = args.output
            if not output_file:
                output_file = "combined_output.md"
                logger.info(f"Output filename not specified, using: {output_file}")

            # Process the directory
            files_by_ext = find_files(directory_path)
            if sum(len(files) for files in files_by_ext.values()) == 0:
                logger.error(f"No supported files found in directory: {directory_path}")
                return 1

            # Rest of directory processing as usual
            # Start generating output
            output_parts = []

            # Add frontmatter
            output_parts.append(generate_yaml_frontmatter(metadata, args.tags))

            # Add title
            output_parts.append("# Combined Document Output\n\n")

            # Process each file type
            all_files = []
            for ext, files in files_by_ext.items():
                for file in files:
                    all_files.append((file, ext))

            # Sort by filename
            all_files.sort(key=lambda x: x[0])

            # Process each file
            for filepath, ext in all_files:
                # Add section header
                output_parts.append(f"## {os.path.basename(filepath)}\n\n")

                # Process file
                content, summary = process_file(
                    filepath, ext, args.summary, args.model, args.length
                )

                # Add summary if generated
                if summary:
                    output_parts.append(summary)

                # Add content
                output_parts.append(content)

                # Add separator
                output_parts.append("---\n\n")

            # Write to output file
            with open(output_file, "w", encoding="utf-8") as f:
                f.write("\n".join(output_parts))

            logger.info(f"Output written to {output_file}")
            return 0
        else:
            logger.error(f"Input path does not exist: {args.input_path}")
            return 1

    # Determine output filename based on input type for flag-based input
    output_file = args.output
    if not output_file:
        if args.file:
            # For a single file, use the same filename but with .md extension
            base_name = os.path.splitext(os.path.basename(args.file))[0]
            output_file = f"{base_name}.md"
        else:
            # For directory or URL, use the default combined output
            output_file = "combined_output.md"
        logger.info(f"Output filename not specified, using: {output_file}")

    # If a URL is provided, process it and exit
    if args.url:
        return process_url_to_markdown(
            args.url, output_file, metadata, args.tags,
            args.summary, args.model, args.length
        )

    # If a single file is provided, process it
    if args.file:
        if not os.path.isfile(args.file):
            logger.error(f"Input file does not exist: {args.file}")
            return 1

        file_path = args.file
        file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')

        # Check if file extension is supported
        if file_ext not in ["pdf", "docx", "txt", "html", "htm", "csv", "md"]:
            logger.error(f"Unsupported file type: {file_ext}")
            return 1

        # Process the single file
        logger.info(f"Processing single file: {file_path}")

        # Start generating output
        output_parts = []

        # Add frontmatter
        output_parts.append(generate_yaml_frontmatter(metadata, args.tags))

        # Add title
        output_parts.append(f"# {os.path.basename(file_path)}\n\n")

        # Process file
        content, summary = process_file(
            file_path, file_ext, args.summary, args.model, args.length
        )

        # Add summary if generated
        if summary:
            output_parts.append(summary)

        # Add content
        output_parts.append(content)

        # Write to output file
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(output_parts))

        logger.info(f"Output written to {output_file}")
        return 0

    # Otherwise, process files from directory
    if args.directory:
        files_by_ext = find_files(args.directory)
        if sum(len(files) for files in files_by_ext.values()) == 0:
            logger.error(f"No supported files found in directory: {args.directory}")
            return 1

        # Start generating output
        output_parts = []

        # Add frontmatter
        output_parts.append(generate_yaml_frontmatter(metadata, args.tags))

        # Add title
        output_parts.append("# Combined Document Output\n\n")

        # Process each file type
        all_files = []
        for ext, files in files_by_ext.items():
            for file in files:
                all_files.append((file, ext))

        # Sort by filename
        all_files.sort(key=lambda x: x[0])

        # Process each file
        for filepath, ext in all_files:
            # Add section header
            output_parts.append(f"## {os.path.basename(filepath)}\n\n")

            # Process file
            content, summary = process_file(
                filepath, ext, args.summary, args.model, args.length
            )

            # Add summary if generated
            if summary:
                output_parts.append(summary)

            # Add content
            output_parts.append(content)

            # Add separator
            output_parts.append("---\n\n")

        # Write to output file
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(output_parts))

        logger.info(f"Output written to {output_file}")
        return 0

    # If we get here, no input source was provided
    logger.error("No input source provided. Please specify a file, directory, or URL.")
    return 1


if __name__ == "__main__":
    sys.exit(main())